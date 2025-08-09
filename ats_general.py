import sys
import os
import re
import spacy
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document
from collections import defaultdict
import easyocr

# Load spaCy
try:
    nlp = spacy.load("en_core_web_sm")
except:
    import subprocess
    subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

# Initialize EasyOCR reader (supports multiple languages)
reader = easyocr.Reader(['en'])

# --- PARAMETERS ---
GENERIC_KEYWORDS = [
    "teamwork", "leadership", "communication", "problem-solving", "initiative",
    "collaboration", "adaptability", "creativity", "critical thinking", "organization"
]
ACTION_VERBS = [
    "achieved", "improved", "trained", "managed", "created", "resolved", "developed", "increased", "researched",
    "led", "designed", "implemented", "launched", "negotiated", "organized", "supervised", "built", "analyzed",
    "delivered", "coordinated", "executed", "optimized", "streamlined", "facilitated", "enhanced"
]
SECTION_PATTERNS = {
    "contact": r"(contact|email|phone|linkedin|address)",
    "summary": r"(summary|objective|profile)",
    "skills": r"(skills|technologies|tools|competencies)",
    "experience": r"(experience|employment|work history|professional experience|career)",
    "education": r"(education|degree|university|college|school)",
    "certifications": r"(certification|certifications|projects|project)"
}
DATE_PATTERN = r"((0[1-9]|1[0-2])\/\d{4}|(Jan(uary)?|Feb(ruary)?|Mar(ch)?|Apr(il)?|May|Jun(e)?|Jul(y)?|Aug(ust)?|Sep(tember)?|Oct(ober)?|Nov(ember)?|Dec(ember)?)\s+\d{4})"

# --- EXTRACT TEXT ---
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = extract_pdf_text(file_path)
    elif ext == ".docx":
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
        # OCR for image files using EasyOCR
        results = reader.readtext(file_path, detail=0)  # detail=0 returns only text
        text = "\n".join(results)
    else:
        raise ValueError("Unsupported file type. Only PDF, DOCX, and image files are supported.")
    # Remove extra spaces, headers/footers (simple heuristic: repeated lines)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    freq = defaultdict(int)
    for line in lines:
        freq[line] += 1
    clean_lines = [line for line in lines if freq[line] < 3]
    return "\n".join(clean_lines)

# --- SECTION PRESENCE ---
def detect_sections(text):
    found = {}
    lower_text = text.lower()
    for section, pattern in SECTION_PATTERNS.items():
        found[section] = bool(re.search(rf"\b{pattern}\b", lower_text))
    # Contact info: look for email, phone, LinkedIn, address
    found['contact'] = bool(re.search(r"\b(email|phone|linkedin|address)\b", lower_text)) or \
        bool(re.search(r"\b\d{10}\b", lower_text)) or \
        bool(re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", lower_text))
    return found

# --- FORMATTING & ATS COMPATIBILITY ---
def check_formatting(file_path, text):
    ext = os.path.splitext(file_path)[1].lower()
    issues = {
        "images": False,
        "tables": False,
        "fonts": True,
        "bullets": False,
        "dates": False
    }
    # Images/tables/fonts
    if ext == ".docx":
        doc = Document(file_path)
        issues["images"] = any(shape for shape in getattr(doc, "inline_shapes", []))
        issues["tables"] = len(doc.tables) > 0
        # Font check: look for more than 3 different fonts
        fonts = set()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    fonts.add(run.font.name)
        issues["fonts"] = len(fonts) <= 3
    elif ext == ".pdf":
        # pdfminer can't check images/tables easily, so flag as unknown
        issues["images"] = "unknown"
        issues["tables"] = "unknown"
        issues["fonts"] = True  # Assume OK
    # Bullets
    issues["bullets"] = bool(re.search(r"^[\-\*\u2022]", text, re.MULTILINE))
    # Dates
    issues["dates"] = bool(re.search(DATE_PATTERN, text, re.IGNORECASE))
    return issues

# --- CONTENT QUALITY ---
def analyze_content(text):
    doc = nlp(text)
    # Action verbs at start of bullet points
    bullet_lines = [line.strip() for line in text.splitlines() if re.match(r"^[\-\*\u2022]", line.strip())]
    action_bullets = 0
    for line in bullet_lines:
        words = line.lstrip("-*•").strip().split()
        if words and words[0].lower() in ACTION_VERBS:
            action_bullets += 1
    # Metrics/numbers
    metrics = len(re.findall(r'\b\d+[\d,\.]*\b', text))
    # Passive voice
    passive_count = sum(1 for sent in doc.sents for token in sent if token.dep_ == "auxpass")
    # Keywords
    keyword_count = sum(1 for kw in GENERIC_KEYWORDS if kw in text.lower())
    return {
        "action_bullets": action_bullets,
        "metrics": metrics,
        "passive_count": passive_count,
        "keyword_count": keyword_count,
        "total_bullets": len(bullet_lines),
        "total_sentences": len(list(doc.sents))
    }

# --- READABILITY & STRUCTURE ---
def check_readability(text):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    # Bullet point length ≤ 2 lines (assume 120 chars per line)
    bullet_lines = [line for line in lines if re.match(r"^[\-\*\u2022]", line)]
    short_bullets = sum(1 for line in bullet_lines if len(line) <= 240)
    # Chronological order: look for descending years in experience
    years = [int(y) for y in re.findall(r'\b(19|20)\d{2}\b', text)]
    chrono = all(earlier >= later for earlier, later in zip(years, years[1:])) if years else True
    # Job titles + company + dates: look for lines with all three (simple heuristic)
    job_blocks = re.findall(r"(?i)(\b(manager|engineer|developer|analyst|consultant|designer|lead|intern|officer|specialist)\b.*?\b(19|20)\d{2}\b)", text)
    job_info = len(job_blocks)
    # Employment gaps > 1 year
    years_sorted = sorted(set(years), reverse=True)
    gaps = any((years_sorted[i] - years_sorted[i+1]) > 1 for i in range(len(years_sorted)-1)) if len(years_sorted) > 1 else False
    return {
        "short_bullets": short_bullets,
        "total_bullets": len(bullet_lines),
        "chrono": chrono,
        "job_info": job_info,
        "gaps": gaps
    }

# --- SCORING ---
def generate_score(sections, formatting, content, readability):
    breakdown = {}
    # Section Presence (30)
    section_score = 0
    section_score += 5 if sections.get("contact") else 0
    section_score += 3 if sections.get("summary") else 0
    section_score += 5 if sections.get("skills") else 0
    section_score += 8 if sections.get("experience") else 0
    section_score += 5 if sections.get("education") else 0
    section_score += 4 if sections.get("certifications") else 0
    breakdown["Section Presence"] = section_score  # max 30

    # Formatting (20)
    formatting_score = 0
    formatting_score += 5 if not formatting["images"] else 0
    formatting_score += 3 if not formatting["tables"] else 0
    formatting_score += 3 if formatting["fonts"] else 0
    formatting_score += 4 if formatting["bullets"] else 0
    formatting_score += 5 if formatting["dates"] else 0
    # If sum > 20, scale down proportionally
    formatting_score = min(formatting_score, 20)
    breakdown["Formatting"] = formatting_score  # max 20

    # Content Quality (30)
    content_score = 0
    content_score += 8 if content["action_bullets"] >= max(1, content["total_bullets"] // 2) else 0
    content_score += 6 if content["metrics"] >= 3 else 0
    content_score += 9 if content["keyword_count"] >= 3 else 0
    content_score += 7 if content["passive_count"] <= (content["total_sentences"] // 5) else 0
    # If sum > 30, scale down proportionally
    content_score = min(content_score, 30)
    breakdown["Content Quality"] = content_score  # max 30

    # Readability & Structure (20)
    read_score = 0
    read_score += 5 if readability["short_bullets"] >= max(1, readability["total_bullets"] // 2) else 0
    read_score += 5 if readability["chrono"] else 0
    read_score += 5 if readability["job_info"] >= 2 else 0
    read_score += 5 if not readability["gaps"] else 0
    breakdown["Readability"] = read_score  # max 20

    total = sum(breakdown.values())
    # Grade
    if total >= 85:
        grade = "Excellent"
    elif total >= 70:
        grade = "Good"
    elif total >= 50:
        grade = "Fair"
    else:
        grade = "Poor"
    return total, breakdown, grade

# --- SUGGESTIONS ---
def generate_suggestions(sections, formatting, content, readability):
    suggestions = []
    # Sections
    if not sections.get("contact"):
        suggestions.append("Add contact information (email, phone, LinkedIn, address).")
    if not sections.get("summary"):
        suggestions.append("Add a Professional Summary or Objective section.")
    if not sections.get("skills"):
        suggestions.append("Add a Skills section with relevant keywords.")
    if not sections.get("experience"):
        suggestions.append("Add a Work Experience section.")
    if not sections.get("education"):
        suggestions.append("Add an Education section.")
    if not sections.get("certifications"):
        suggestions.append("Include Certifications or Projects if relevant.")
    # Formatting
    if formatting["images"]:
        suggestions.append("Avoid using images or charts—ATS may not parse them.")
    if formatting["tables"]:
        suggestions.append("Avoid using tables—ATS may not parse them.")
    if not formatting["fonts"]:
        suggestions.append("Use consistent fonts and sizes throughout your resume.")
    if not formatting["bullets"]:
        suggestions.append("Use bullet points for experiences instead of paragraphs.")
    if not formatting["dates"]:
        suggestions.append("Use proper date formats (MM/YYYY or Month YYYY) for jobs and education.")
    # Content
    if content["action_bullets"] < max(1, content["total_bullets"] // 2):
        suggestions.append("Start more bullet points with strong action verbs.")
    if content["metrics"] < 3:
        suggestions.append("Add measurable results or numbers to your achievements.")
    if content["passive_count"] > (content["total_sentences"] // 5):
        suggestions.append("Reduce passive voice; use active voice for stronger statements.")
    if content["keyword_count"] < 3:
        suggestions.append("Include more professional keywords (teamwork, leadership, etc.).")
    # Readability
    if readability["short_bullets"] < max(1, readability["total_bullets"] // 2):
        suggestions.append("Keep bullet points concise (≤ 2 lines).")
    if not readability["chrono"]:
        suggestions.append("List work experience in clear, reverse-chronological order.")
    if readability["job_info"] < 2:
        suggestions.append("Ensure job titles, company names, and dates are present for each job.")
    if readability["gaps"]:
        suggestions.append("Explain any employment gaps longer than 1 year.")
    if not suggestions:
        suggestions.append("Your resume is ATS-friendly!")
    return suggestions

# --- OUTPUT TABLE ---
def print_table(score, breakdown, grade, suggestions):
    print("="*50)
    print(f"ATS-Friendliness Score: {score}/100   Grade: {grade}")
    print("="*50)
    print("{:<25} {:>10}/30".format("Section Presence", breakdown["Section Presence"]))
    print("{:<25} {:>10}/20".format("Formatting", breakdown["Formatting"]))
    print("{:<25} {:>10}/30".format("Content Quality", breakdown["Content Quality"]))
    print("{:<25} {:>10}/20".format("Readability", breakdown["Readability"]))
    print("="*50)
    print("Suggestions:")
    for s in suggestions:
        print(f"- {s}")
    print("="*50)

# --- MAIN ---
def main():
    if len(sys.argv) != 2:
        print("Usage: python ats.py <resume.pdf|resume.docx>")
        sys.exit(1)
    file_path = sys.argv[1]
    try:
        text = extract_text(file_path)
        sections = detect_sections(text)
        formatting = check_formatting(file_path, text)
        content = analyze_content(text)
        readability = check_readability(text)
        score, breakdown, grade = generate_score(sections, formatting, content, readability)
        suggestions = generate_suggestions(sections, formatting, content, readability)
        print_table(score, breakdown, grade, suggestions)
    except Exception as e:
        print(f"Error: {e}")

def analyze_resume(file_path):
    """
    Analyze a resume file and return ATS score, breakdown, grade, and suggestions.
    Args:
        file_path (str): Path to the resume PDF or DOCX file.
    Returns:
        dict: {
            "score": int,
            "breakdown": dict,
            "grade": str,
            "suggestions": list of str
        }
    """
    text = extract_text(file_path)
    sections = detect_sections(text)
    formatting = check_formatting(file_path, text)
    content = analyze_content(text)
    readability = check_readability(text)
    score, breakdown, grade = generate_score(sections, formatting, content, readability)
    suggestions = generate_suggestions(sections, formatting, content, readability)
    return {
        "score": score,
        "breakdown": breakdown,
        "grade": grade,
        "suggestions": suggestions
    }

# Keep the CLI main for direct execution
if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python ats.py <resume.pdf|resume.docx>")
        sys.exit(1)
    file_path = sys.argv[1]
    try:
        result = analyze_resume(file_path)
        print_table(result["score"], result["breakdown"], result["grade"], result["suggestions"])
    except Exception as e:
        print(f"Error: {e}")